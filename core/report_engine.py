"""
core/report_engine.py – Report generation in PDF, HTML, Markdown, JSON, CSV, DOCX.
"""

from __future__ import annotations

import csv
import json
import logging
from datetime import datetime, timezone
from io import StringIO
from pathlib import Path
from typing import Any, Optional, TYPE_CHECKING

if TYPE_CHECKING:
    from core.database import Database
    from core.config import Config

log = logging.getLogger(__name__)


class ReportEngine:
    """Generate investigation reports in multiple formats."""

    def __init__(self, config: "Config", db: "Database") -> None:
        self._config = config
        self._db = db

    # ── Public API ─────────────────────────────────────────────────────────

    def generate(
        self,
        case_id: int,
        fmt: str,
        output_path: Optional[Path] = None,
        title: Optional[str] = None,
        include_ai_summary: str = "",
    ) -> Path:
        """
        Generate a report for a case.

        Args:
            case_id:          Case to report on.
            fmt:              One of pdf|html|markdown|json|csv|docx.
            output_path:      Override output file path.
            title:            Override report title.
            include_ai_summary: Pre-generated AI narrative to embed.

        Returns:
            Path to the generated report file.
        """
        case = self._db.fetchone("SELECT * FROM cases WHERE id=?", (case_id,))
        if not case:
            raise ValueError(f"Case {case_id} not found")

        report_title = title or f"SentinelAI Report – {case['title']}"
        out_dir = Path(
            self._config.get("report_output_dir", str(Path.home() / "SentinelAI_Reports"))
        )
        out_dir.mkdir(parents=True, exist_ok=True)

        ts = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
        safe_title = "".join(c if c.isalnum() or c in " _-" else "_" for c in case["title"])
        fmt = fmt.lower()
        ext_map = {
            "pdf": ".pdf", "html": ".html", "markdown": ".md",
            "json": ".json", "csv": ".csv", "docx": ".docx",
        }
        if output_path is None:
            output_path = out_dir / f"{safe_title}_{ts}{ext_map.get(fmt, '.txt')}"

        data = self._collect_data(case_id, dict(case), include_ai_summary, report_title)

        dispatch = {
            "pdf": self._write_pdf,
            "html": self._write_html,
            "markdown": self._write_markdown,
            "json": self._write_json,
            "csv": self._write_csv,
            "docx": self._write_docx,
        }
        writer = dispatch.get(fmt, self._write_markdown)
        writer(data, output_path)

        # Record in DB
        self._db.insert(
            "INSERT INTO reports (case_id, title, format, file_path) VALUES (?,?,?,?)",
            (case_id, report_title, fmt, str(output_path)),
        )
        log.info("Report generated: %s", output_path)
        return output_path

    # ── Data collection ────────────────────────────────────────────────────

    def _collect_data(
        self, case_id: int, case: dict, ai_summary: str, title: str
    ) -> dict:
        evidence = [dict(r) for r in self._db.fetchall(
            "SELECT * FROM evidence WHERE case_id=? ORDER BY created_at", (case_id,)
        )]
        iocs = [dict(r) for r in self._db.fetchall(
            "SELECT * FROM iocs WHERE case_id=? ORDER BY ioc_type, value", (case_id,)
        )]
        timeline = [dict(r) for r in self._db.fetchall(
            """SELECT * FROM timeline_entries WHERE case_id=?
               ORDER BY timestamp""", (case_id,)
        )]
        _ec = self._db.fetchone(
            "SELECT COUNT(*) AS c FROM log_events WHERE case_id=?", (case_id,)
        )
        events_count = _ec["c"] if _ec else 0
        _fc = self._db.fetchone(
            "SELECT COUNT(*) AS c FROM log_events WHERE case_id=? AND flagged=1",
            (case_id,)
        )
        flagged_count = _fc["c"] if _fc else 0
        return {
            "title": title,
            "generated_at": datetime.now(timezone.utc).isoformat(),
            "case": case,
            "evidence": evidence,
            "iocs": iocs,
            "timeline": timeline,
            "events_count": events_count,
            "flagged_count": flagged_count,
            "ai_summary": ai_summary,
        }

    # ── PDF ────────────────────────────────────────────────────────────────

    def _write_pdf(self, data: dict, path: Path) -> None:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable,
        )

        doc = SimpleDocTemplate(str(path), pagesize=A4,
                                leftMargin=2*cm, rightMargin=2*cm,
                                topMargin=2*cm, bottomMargin=2*cm)
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle("Title", parent=styles["Title"],
                                     textColor=colors.HexColor("#0D47A1"), fontSize=18)
        h2_style = ParagraphStyle("H2", parent=styles["Heading2"],
                                  textColor=colors.HexColor("#1565C0"), fontSize=13)
        body = styles["Normal"]
        story = []

        story.append(Paragraph(data["title"], title_style))
        story.append(Spacer(1, 0.3*cm))
        story.append(Paragraph(
            f"Generated: {data['generated_at']} | Case ID: {data['case']['id']}",
            body,
        ))
        story.append(HRFlowable(width="100%", color=colors.HexColor("#1565C0")))
        story.append(Spacer(1, 0.4*cm))

        # Case summary
        story.append(Paragraph("Case Summary", h2_style))
        c = data["case"]
        summary_rows = [
            ["Title", c.get("title", "")],
            ["Severity", c.get("severity", "").upper()],
            ["Status", c.get("status", "").capitalize()],
            ["Analyst", c.get("analyst", "")],
            ["Tags", c.get("tags", "")],
            ["Description", c.get("description", "")],
        ]
        t = Table(summary_rows, colWidths=[4*cm, 13*cm])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#E3F2FD")),
            ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        story.append(t)
        story.append(Spacer(1, 0.4*cm))

        # Statistics
        story.append(Paragraph("Statistics", h2_style))
        stat_rows = [
            ["Total Log Events", str(data["events_count"])],
            ["Flagged Events", str(data["flagged_count"])],
            ["Evidence Files", str(len(data["evidence"]))],
            ["IOCs", str(len(data["iocs"]))],
            ["Timeline Entries", str(len(data["timeline"]))],
        ]
        t2 = Table(stat_rows, colWidths=[6*cm, 11*cm])
        t2.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#E8F5E9")),
            ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ]))
        story.append(t2)
        story.append(Spacer(1, 0.4*cm))

        # AI Summary
        if data["ai_summary"]:
            story.append(Paragraph("AI Investigation Summary", h2_style))
            story.append(Paragraph(data["ai_summary"].replace("\n", "<br/>"), body))
            story.append(Spacer(1, 0.4*cm))

        # IOCs
        if data["iocs"]:
            story.append(Paragraph("Indicators of Compromise", h2_style))
            ioc_rows = [["Type", "Value", "Confidence", "Threat Type", "Notes"]]
            for ioc in data["iocs"]:
                ioc_rows.append([
                    ioc.get("ioc_type", ""), ioc.get("value", "")[:60],
                    ioc.get("confidence", ""), ioc.get("threat_type", ""),
                    ioc.get("notes", "")[:40],
                ])
            t3 = Table(ioc_rows, colWidths=[2.5*cm, 5*cm, 2.5*cm, 3*cm, 4*cm])
            t3.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1565C0")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F5F5F5")]),
            ]))
            story.append(t3)
            story.append(Spacer(1, 0.4*cm))

        # Timeline
        if data["timeline"]:
            story.append(Paragraph("Timeline", h2_style))
            for entry in data["timeline"][:100]:
                story.append(Paragraph(
                    f"<b>[{entry.get('timestamp', '')}]</b> "
                    f"<i>{entry.get('event_type', '')}</i> – "
                    f"{entry.get('description', '')}",
                    body,
                ))
            story.append(Spacer(1, 0.2*cm))

        doc.build(story)

    # ── HTML ───────────────────────────────────────────────────────────────

    def _write_html(self, data: dict, path: Path) -> None:
        html = self._render_html(data)
        path.write_text(html, encoding="utf-8")

    def _render_html(self, data: dict) -> str:
        c = data["case"]
        ioc_rows = "".join(
            f"<tr><td>{i.get('ioc_type','')}</td><td>{i.get('value','')}</td>"
            f"<td>{i.get('confidence','')}</td><td>{i.get('threat_type','')}</td></tr>"
            for i in data["iocs"]
        )
        tl_rows = "".join(
            f"<tr><td>{e.get('timestamp','')}</td><td>{e.get('event_type','')}</td>"
            f"<td>{e.get('description','')}</td><td>{e.get('mitre_tech','')}</td></tr>"
            for e in data["timeline"]
        )
        ai_block = f"<section><h2>AI Summary</h2><pre>{data['ai_summary']}</pre></section>" \
                   if data["ai_summary"] else ""
        return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>{data['title']}</title>
<style>
body{{font-family:'Segoe UI',Arial,sans-serif;background:#0D1117;color:#C9D1D9;margin:0;padding:2em}}
h1{{color:#58A6FF;border-bottom:2px solid #21262D}}
h2{{color:#79C0FF;margin-top:2em}}
table{{width:100%;border-collapse:collapse;margin:1em 0}}
th{{background:#161B22;color:#58A6FF;padding:.6em;text-align:left}}
td{{padding:.5em;border-bottom:1px solid #21262D}}
tr:nth-child(even){{background:#161B22}}
.stat{{display:inline-block;background:#161B22;border:1px solid #30363D;
        border-radius:6px;padding:1em 2em;margin:.5em;text-align:center}}
.stat .num{{font-size:2em;color:#58A6FF;font-weight:700}}
pre{{background:#161B22;padding:1em;border-radius:4px;white-space:pre-wrap;overflow-x:auto}}
</style>
</head>
<body>
<h1>{data['title']}</h1>
<p>Generated: {data['generated_at']} | Case ID: {c['id']} | Analyst: {c.get('analyst','')}</p>
<div>
  <div class="stat"><div class="num">{data['events_count']}</div>Log Events</div>
  <div class="stat"><div class="num">{data['flagged_count']}</div>Flagged</div>
  <div class="stat"><div class="num">{len(data['iocs'])}</div>IOCs</div>
  <div class="stat"><div class="num">{len(data['timeline'])}</div>Timeline</div>
</div>
<section>
<h2>Case Details</h2>
<table><tr><th>Field</th><th>Value</th></tr>
<tr><td>Title</td><td>{c.get('title','')}</td></tr>
<tr><td>Severity</td><td>{c.get('severity','').upper()}</td></tr>
<tr><td>Status</td><td>{c.get('status','')}</td></tr>
<tr><td>Description</td><td>{c.get('description','')}</td></tr>
</table>
</section>
{ai_block}
<section>
<h2>Indicators of Compromise ({len(data['iocs'])})</h2>
<table><tr><th>Type</th><th>Value</th><th>Confidence</th><th>Threat Type</th></tr>
{ioc_rows}
</table>
</section>
<section>
<h2>Timeline ({len(data['timeline'])} entries)</h2>
<table><tr><th>Timestamp</th><th>Event Type</th><th>Description</th><th>MITRE</th></tr>
{tl_rows}
</table>
</section>
</body></html>"""

    # ── Markdown ───────────────────────────────────────────────────────────

    def _write_markdown(self, data: dict, path: Path) -> None:
        c = data["case"]
        lines = [
            f"# {data['title']}",
            f"",
            f"**Generated:** {data['generated_at']}  ",
            f"**Case ID:** {c['id']}  ",
            f"**Severity:** {c.get('severity','').upper()}  ",
            f"**Analyst:** {c.get('analyst','')}",
            "",
            "## Summary Statistics",
            "",
            f"| Metric | Value |",
            f"|--------|-------|",
            f"| Log Events | {data['events_count']} |",
            f"| Flagged Events | {data['flagged_count']} |",
            f"| Evidence Files | {len(data['evidence'])} |",
            f"| IOCs | {len(data['iocs'])} |",
            f"| Timeline Entries | {len(data['timeline'])} |",
            "",
        ]
        if data["ai_summary"]:
            lines += ["## AI Investigation Summary", "", data["ai_summary"], ""]

        if data["iocs"]:
            lines += ["## Indicators of Compromise", "",
                      "| Type | Value | Confidence | Threat |",
                      "|------|-------|------------|--------|"]
            for i in data["iocs"]:
                lines.append(
                    f"| {i.get('ioc_type','')} | `{i.get('value','')}` "
                    f"| {i.get('confidence','')} | {i.get('threat_type','')} |"
                )
            lines.append("")

        if data["timeline"]:
            lines += ["## Timeline", "",
                      "| Timestamp | Event Type | Description | MITRE |",
                      "|-----------|------------|-------------|-------|"]
            for e in data["timeline"]:
                lines.append(
                    f"| {e.get('timestamp','')} | {e.get('event_type','')} "
                    f"| {e.get('description','')} | {e.get('mitre_tech','')} |"
                )
        path.write_text("\n".join(lines), encoding="utf-8")

    # ── JSON ───────────────────────────────────────────────────────────────

    def _write_json(self, data: dict, path: Path) -> None:
        path.write_text(json.dumps(data, indent=2, default=str), encoding="utf-8")

    # ── CSV ────────────────────────────────────────────────────────────────

    def _write_csv(self, data: dict, path: Path) -> None:
        with path.open("w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow(["Section", "Field", "Value"])
            c = data["case"]
            for k, v in c.items():
                writer.writerow(["Case", k, v])
            writer.writerow([])
            writer.writerow(["IOC Type", "Value", "Confidence", "Threat Type", "Notes"])
            for ioc in data["iocs"]:
                writer.writerow([
                    ioc.get("ioc_type",""), ioc.get("value",""),
                    ioc.get("confidence",""), ioc.get("threat_type",""),
                    ioc.get("notes",""),
                ])
            writer.writerow([])
            writer.writerow(["Timestamp", "Event Type", "Description", "MITRE Tech", "Severity"])
            for e in data["timeline"]:
                writer.writerow([
                    e.get("timestamp",""), e.get("event_type",""),
                    e.get("description",""), e.get("mitre_tech",""),
                    e.get("severity",""),
                ])

    # ── DOCX ───────────────────────────────────────────────────────────────

    def _write_docx(self, data: dict, path: Path) -> None:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        c = data["case"]

        # Title
        title_para = doc.add_heading(data["title"], 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(
            f"Generated: {data['generated_at']} | Case ID: {c['id']}"
        )
        doc.add_heading("Case Summary", 1)
        t = doc.add_table(rows=0, cols=2)
        t.style = "Table Grid"
        for field, value in [
            ("Title", c.get("title", "")),
            ("Severity", c.get("severity", "").upper()),
            ("Status", c.get("status", "")),
            ("Analyst", c.get("analyst", "")),
            ("Description", c.get("description", "")),
        ]:
            row = t.add_row().cells
            row[0].text = field
            row[1].text = str(value)

        doc.add_heading("Statistics", 1)
        t2 = doc.add_table(rows=0, cols=2)
        t2.style = "Table Grid"
        for label, val in [
            ("Log Events", data["events_count"]),
            ("Flagged", data["flagged_count"]),
            ("Evidence Files", len(data["evidence"])),
            ("IOCs", len(data["iocs"])),
            ("Timeline Entries", len(data["timeline"])),
        ]:
            r = t2.add_row().cells
            r[0].text = label
            r[1].text = str(val)

        if data["ai_summary"]:
            doc.add_heading("AI Investigation Summary", 1)
            doc.add_paragraph(data["ai_summary"])

        if data["iocs"]:
            doc.add_heading(f"Indicators of Compromise ({len(data['iocs'])})", 1)
            t3 = doc.add_table(rows=1, cols=4)
            t3.style = "Table Grid"
            hdr = t3.rows[0].cells
            hdr[0].text = "Type"
            hdr[1].text = "Value"
            hdr[2].text = "Confidence"
            hdr[3].text = "Threat Type"
            for ioc in data["iocs"]:
                r = t3.add_row().cells
                r[0].text = ioc.get("ioc_type", "")
                r[1].text = ioc.get("value", "")
                r[2].text = ioc.get("confidence", "")
                r[3].text = ioc.get("threat_type", "")

        if data["timeline"]:
            doc.add_heading("Timeline", 1)
            t4 = doc.add_table(rows=1, cols=4)
            t4.style = "Table Grid"
            hdr2 = t4.rows[0].cells
            hdr2[0].text = "Timestamp"
            hdr2[1].text = "Event"
            hdr2[2].text = "Description"
            hdr2[3].text = "MITRE"
            for e in data["timeline"]:
                r = t4.add_row().cells
                r[0].text = e.get("timestamp", "")
                r[1].text = e.get("event_type", "")
                r[2].text = e.get("description", "")
                r[3].text = e.get("mitre_tech", "")

        doc.save(str(path))
